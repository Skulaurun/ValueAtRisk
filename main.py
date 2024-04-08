import os
import json
import urllib.request
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from io import BytesIO
import scipy.stats
import numpy as np
import pandas as pd
import yfinance as yf
import plotly.graph_objs as go
from plotly.io import to_image
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION

# <Configuration>

# VaR is calculated every month, how many VaRs to calculate,
# default is 24
BACKTEST_MONTHS = 24

# How much historic data should be used for VaR calculation,
# default is 5 years
BACKTEST_MAX_YEARS = 5

# The backtest needs at least 2 years worth of historic data
BACKTEST_MIN_YEARS = 2

# Monthly VaR
MONTE_CARLO_TIME_SPAN = 22 # trading days

# How many Monte Carlo simulations to run
MONTE_CARLO_NUM_SIMULATIONS = 10000

# Directory where to output reports,
# created automatically if doesn't exist
OUTPUT_DIRECTORY = './reports'

# </Configuration>

# Plotly doesn't add shapes added using `add_shape` to the legend,
# so we create an invisible trace
def add_line_with_legend(fig, x0, x1, y0, y1, line, description):
    fig.add_shape(
        type='line',
        x0=x0,
        x1=x1,
        y0=y0,
        y1=y1,
        line=line
    )
    fig.add_trace(go.Scatter(
        x=[None],
        y=[None],
        mode='lines',
        line=line,
        name=description
    ))

# Generate graph for three VaRs and price
# : info = yfinance ticker.info
# : close_prices = pandas dataframe of close prices acquired using yfinance
# : VaRs = array of tuples [(confidence, color, value)]
def generate_graph(info, close_prices, VaRs=[]):
    # Since the VaR line needs to be added after the last price,
    # we extend the x axis (time) by one month

    # Get the last date in the current data
    last_date = close_prices.index[-1]
    # Calculate the first date of the next month
    next_month_date = last_date + pd.DateOffset(months=1)
    # Create a new datetime index for the additional month
    new_dates = pd.date_range(start=next_month_date, periods=20, freq='B')
    # Create a new DataFrame with NaN values for the additional month
    additional_month_data = pd.DataFrame(index=new_dates, columns=['Close'])
    # Concatenate the existing data with the new data
    extended_close_prices = pd.concat([close_prices, additional_month_data['Close']])

    # Real last price
    last_price_x0 = close_prices.index[-1]
    # Fictional last price
    last_price_x1 = extended_close_prices.index[-1]

    fig = go.Figure(data=go.Scatter(x=extended_close_prices.index, y=extended_close_prices, name='Price'))
    # Add graph and axis titles
    fig.update_layout(title=f'{info.get("shortName")} ({info.get("symbol")})', xaxis_title='Datum', yaxis_title='Cena')

    # VaR[0] - confidence 0 - 1
    # VaR[1] - color
    # VaR[2] - value
    for VaR in VaRs:
        add_line_with_legend(
            fig,
            x0=last_price_x0,
            x1=last_price_x1,
            y0=VaR[2],
            y1=VaR[2],
            line=dict(color=VaR[1], width=2, dash='dash'),
            description=f'VaR {int(VaR[0] * 100)}% (={round(VaR[2], 2)})'
        )

    # Add legend (and set its position)
    fig.update_layout(legend=dict(x=0.74, y=1.3, bgcolor='rgba(255, 255, 255, 0.5)'))
    return fig

# Calculate VaR
# : ticker = yfinance ticker
# : method = function, one of (historic_var, varcov_var, monte_carlo_var)
# : confidences = array of tuples [(confidence, color)]
def calculate_var(ticker, method, confidences=[(0.99, 'red')]):
    data = ticker.history(period='12mo')
    close_prices = data['Close']
    VaRs = []
    for confidence in confidences:
        value_at_risk = method(close_prices, confidence[0])
        # The tuple is extended to (confidence, color, last_price + VaR)
        # Value at Risk is subtracted from last price (VaR is negative number)
        VaRs.append(confidence + (close_prices.iloc[-1] + value_at_risk,))
    return generate_graph(ticker.info, close_prices, VaRs)

# Calculate returns (in absolute values)
# Take the next close price and subtract the previous
def calculate_returns(close_prices):
    returns = []
    for i in range(len(close_prices) - 1):
        returns.append(close_prices.iloc[i + 1] - close_prices.iloc[i])
    return returns

# Calculate VaR solely based on historical data
def historic_var(close_prices, confidence):
    returns = calculate_returns(close_prices)
    return np.percentile(returns, 100 * (1 - confidence))

# Monte Carlo VaR
def monte_carlo_var(close_prices, confidence):
    returns = calculate_returns(close_prices)

    average_daily_return = np.mean(returns) # mu
    standard_deviation = np.std(returns) # sigma

    def random_z_score():
        return np.random.normal(0, 1)

    # Equation: R = mu * dt + sigma * random * sqrt(dt)
    # Source: https://financetrain.com/calculating-var-using-monte-carlo-simulation
    def scenario_gain_loss(standard_deviation, z_score, days):
        return (average_daily_return * days) \
            + (standard_deviation * z_score * np.sqrt(days))

    # Perform Monte-Carlo simulations
    scenario_returns = []
    for i in range(MONTE_CARLO_NUM_SIMULATIONS):
        z_score = random_z_score() # Include randomness in the calculation
        scenario_returns.append(scenario_gain_loss(standard_deviation, z_score, MONTE_CARLO_TIME_SPAN))

    return np.percentile(scenario_returns, 100 * (1 - confidence))

# Variance-Covariance VaR
def varcov_var(close_prices, confidence):
    # Source: https://github.com/MBKraus/Python_Portfolio__VaR_Tool/blob/master/Portfolio_VaR_Toolv5.py
    returns = calculate_returns(close_prices)
    average_daily_return = np.mean(returns) # mu
    standard_deviation = np.std(returns) # sigma
    # This if-block was added to prevent an error being thrown
    if average_daily_return == 0 and standard_deviation == 0:
        return 0 # price is equal => no returns
    return scipy.stats.norm.ppf(1 - confidence, average_daily_return, standard_deviation)

# Add plotly figure to Word document
# : document = python-docx document
# : graph = plotly Figure
def add_graph(document, graph):
    # Convert plotly figure to byte buffer
    image_buffer = BytesIO(to_image(graph, format='png'))
    # Add picture to document from memory buffer
    # The picture has fixed size of 16.5cm
    document.add_picture(image_buffer, width=Cm(16.5))

# Insert simple dictionary as a table
# : document = python-docx document
# : info = dictionary of { 'Row name': value }
def insert_info(document, info):
    i = 0
    table = document.add_table(rows=len(info.items()), cols=2)
    for key, value in info.items():
        row = table.rows[i].cells
        row[0].text = str(key)
        row[1].text = str(value)
        i += 1

# Insert pandas dataframe as a table (fetched from yfinance)
# This function allows field filtering
# : document = python-docx document
# : balance_sheet = pandas dataframe (doesn't need to be Balance sheet)
# : wanted_fields = array of strings ['Total Revenue', 'Cost of Revenue']
def insert_table_filtered(document, balance_sheet, wanted_fields):
    # We create a copy of the dataframe (because we will modify it inside the loop)
    for field in balance_sheet.copy().index:
        if not field in wanted_fields:
            # Drop unwanted fields
            balance_sheet.drop(index = field, inplace=True)
    # Create table and set its style
    table = document.add_table(rows=len(balance_sheet.index) + 1, cols=len(balance_sheet.columns) + 1)
    table.style = 'Light Grid'
    # Add header row
    for i in range(len(balance_sheet.columns)):
        # Format date and insert inside the first row of the table (the header row)
        table.cell(0, i + 1).text = datetime.strftime(balance_sheet.columns[i], '%m/%d/%Y')
    # Add fields
    for i, field in enumerate(balance_sheet.index):
        # Field name, e.g. Net Income Common Stockholders
        table.cell(i + 1, 0).text = field
        # Fill values into the table as time series
        for j in range(len(balance_sheet.columns)):
            value = balance_sheet.loc[field].iloc[j]
            value /= 1000 # Convert to thousands
            # Convert NaN values to N/A
            table.cell(i + 1, j + 1).text = "{:,}".format(int(value)) if not pd.isna(value) else "N/A"
    # Add padding/space before and after the table so it looks good
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)

# Insert pandas dataframe as a table (fetched from yfinance)
# : document = python-docx document
# : income_statement = pandas dataframe (doesn't need to be Income statement)
# : to_thousands = boolean if True converts values to thousands using division
def insert_table(document, income_statement, to_thousands = True):
    # Create table and set its style
    table = document.add_table(rows=len(income_statement.index) + 1, cols=len(income_statement.columns) + 1)
    table.style = 'Light Grid'
    # Add header row
    for i in range(len(income_statement.columns)):
        # Format date and insert inside the first row of the table (the header row)
        table.cell(0, i + 1).text = datetime.strftime(income_statement.columns[i], '%m/%d/%Y')
    # Add fields
    for i, field in enumerate(income_statement.index):
        # Field name, e.g. Net Income Common Stockholders
        table.cell(i + 1, 0).text = field
        # Fill values into the table as time series
        for j in range(len(income_statement.columns)):
            value = income_statement.loc[field].iloc[j]
            # Convert to thousand if specified by a function parameter
            if to_thousands:
                value /= 1000
            # This value is the output field value (not necessarily a string yet)
            string_value = value
            # To round we need a non NaN int
            if not pd.isna(value):
                # To round we need an int
                string_value = int(value)
                if not to_thousands:
                    # Round to 2 decimal places
                    string_value = round(value, 2)
            # Convert NaN output values to N/A
            table.cell(i + 1, j + 1).text = "{:,}".format(string_value) if not pd.isna(value) else "N/A"
    # Add padding/space before and after the table so it looks good
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)

# Fetch ESG chart using direct request to Yahoo Finance API,
# as at the time of writing this script, yfinance library doesn't support it natively
# : document = python-docx document
# : symbol = Yahoo Finance ticker symbol (a string)
def try_insert_sustainability(document, symbol):
    try:
        # Source: https://stackoverflow.com/a/67388044
        url = f'https://query2.finance.yahoo.com/v1/finance/esgChart?symbol={symbol}'
        response = urllib.request.urlopen(url)
        data = json.loads(response.read())
        # Convert the dataframe to the same format yfinance library uses,
        # so we can use the same functions for inserting into the document,
        # e.g. an insert_table() or insert_table_filtered()
        df = pd.DataFrame(data["esgChart"]["result"][0]["symbolSeries"])
        df.index = pd.to_datetime(df["timestamp"], unit="s")
        df = df.drop("timestamp", axis=1)
        df = df.dropna().tail().transpose()
        row_names = {
            'esgScore': 'Total ESG Risk Score',
            'governanceScore': 'Governance Risk Score',
            'environmentScore': 'Environment Risk Score',
            'socialScore': 'Social Risk Score'
        }
        df = df.rename(index = row_names)
        df = df.iloc[:, ::-1] # reverse dataframe
        heading_income = document.add_heading('Environment, Social and Governance (ESG)', 1)
        heading_income.paragraph_format.space_before = Pt(4)
        insert_table(document, df, False)
        document.add_section(WD_SECTION.NEW_PAGE)
    except:
        pass

# Generate backtest VaR lines (horizontal and vertical)
# : fig = plotly figure
# : dates = x axis (time)
# : values = calculated VaRs (already subtracted from last price)
# : color = line color, e.g. 'red'
def make_line_for_backtest(fig, dates, values, color):
    for i in range(0, len(dates) - 1):
        # Horizontal
        fig.add_shape(
            type='line',
            x0=dates[i],
            x1=dates[i + 1],
            y0=values[i],
            y1=values[i],
            line=dict(color=color, width=2)
        )
        # Vertical
        fig.add_shape(
            type='line',
            x0=dates[i + 1],
            x1=dates[i + 1],
            y0=values[i],
            y1=values[i + 1],
            line=dict(color=color, width=1, dash='dash')
        )

# Perform bakctesting
# : ticker = yfinance ticker
# : method = function, one of (historic_var, varcov_var, monte_carlo_var)
# : confidences = array of tuples [(confidence, color)]
def do_backtest(ticker, method, confidences):
    fig = go.Figure()

    # Retrieve historic data for close prices,
    # we need only so much data equal to how many VaRs we want to show
    #end_date = datetime.now().replace(day=1)
    end_date = datetime.now()
    start_date = end_date - relativedelta(months=BACKTEST_MONTHS)
    history = ticker.history(start=start_date, end=end_date)

    # Since the VaR line needs to be added after the last price,
    # we extend the x axis (time) by one month
    close_prices_price_line = history['Close']
    # Get the last date in the current data
    last_date = close_prices_price_line.index[-1]
    # Calculate the first date of the next month
    next_month_date = last_date + pd.DateOffset(months=1)
    # Create a new datetime index for the additional month
    new_dates = pd.date_range(start=next_month_date, periods=20, freq='B')
    # Create a new DataFrame with NaN values for the additional month
    additional_month_data = pd.DataFrame(index=new_dates, columns=['Close'])
    # Concatenate the existing data with the new data
    extended_close_prices = pd.concat([close_prices_price_line, additional_month_data['Close']])

    price_trace = go.Scatter(x=extended_close_prices.index, y=extended_close_prices, mode='lines', name='Price')

    fig.add_trace(price_trace)
    # Add graph and axis titles
    fig.update_layout(title=f'{ticker.info.get("shortName")} ({ticker.info.get("symbol")})',
                    xaxis_title='Datum',
                    yaxis_title='Cena')

    # Perform backtesting for each confidence value
    for confidence in confidences:
        VaRs = []
        dates = []
        prices = []
        isFirstRun = True
        for i in range(0, BACKTEST_MONTHS + 1):
            #end_date = datetime.now().replace(day=1) - relativedelta(months=i)
            # Calculate fictional present date (end)
            end_date = datetime.now() - relativedelta(months=i)
            # Calculate past date (start), how many years worth of data to use
            start_date = end_date - timedelta(days=365 * BACKTEST_MAX_YEARS)

            try:
                # Fetch data from yfinance for time frame
                data = ticker.history(start=start_date, end=end_date)
                try:
                    # Calculate how much data we have in years (it is enough?)
                    difference = end_date.timestamp() - pd.to_datetime(data.index.min()).timestamp()
                    difference_years = int(round(difference / (60 * 60 * 24 * 365.25), 0))
                    # If first run - calculate at least one VaR from the data we have
                    if not isFirstRun and difference_years < BACKTEST_MIN_YEARS:
                        print(f"Not enough data, (difference_years = {difference_years}) < {BACKTEST_MIN_YEARS}")
                        print("The backtest won't go any further into the past.\n")
                        break
                except:
                    # NaN error (not enough data in pandas dataframe to calculate difference)
                    print(f"Error during backtesting [ method = {method.__name__} ], historical data not available for given time period.")
                    print("The backtest won't go any further into the past.\n")
                    break
            except:
                # Error from yfinance library (couldn't fetch the data for specific time period)
                print(f"Error during backtesting [ method = {method.__name__} ], historical data not available for given time period.")
                print("The backtest won't go any further into the past.\n")
                break

            close_prices = data['Close']
            if close_prices.empty:
                print(f"Error during backtesting [ method = {method.__name__} ], historical data not available for given time period.")
                print("The backtest won't go any further into the past.\n")
                break

            dates.append(end_date)
            prices.append(close_prices.iloc[-1])
            VaRs.append(close_prices.iloc[-1] + method(close_prices, confidence[0]))
            isFirstRun = False

        # Reverse arrays (since items are added from present to past)
        dates = dates[::-1]
        VaRs = VaRs[::-1]
        prices = prices[::-1]

        # Duplicate last (for final/future horizontal line)
        dates.append(datetime.now() + relativedelta(months=1))
        if len(VaRs) >= 1:
            VaRs.append(VaRs[-1])

        # Display last VaR calculated in the graph legend (as a value)
        last_var_to_display = np.nan
        if len(VaRs) >= 1:
            last_var_to_display = round(VaRs[-1], 2)

        # Graph all VaRs
        make_line_for_backtest(fig, dates, VaRs, confidence[1])
        fig.add_trace(go.Scatter(
            x=[None],
            y=[None],
            mode='lines',
            line=dict(color=confidence[1], width=2, dash='dash'),
            name=f'VaR {int(confidence[0] * 100)}% (={last_var_to_display})'
        ))

    fig.update_layout(legend=dict(x=0.74, y=1.3, bgcolor='rgba(255, 255, 255, 0.5)'))
    return fig

# Create report for ticker/stock
# : ticker = yfinance ticker
# : output_path = file path where to save word document
def create_report(ticker, output_path):
    CONFIDENCES = [
        (0.99, 'red'),
        (0.97, 'magenta'),
        (0.95, 'green')
    ]
    historic_graph = calculate_var(ticker, method=historic_var, confidences=CONFIDENCES)
    varcov_graph = calculate_var(ticker, method=varcov_var, confidences=CONFIDENCES)
    monte_carlo_graph = calculate_var(ticker, method=monte_carlo_var, confidences=CONFIDENCES)

    historic_backtest = do_backtest(ticker, method=historic_var, confidences=CONFIDENCES)
    varcov_backtest = do_backtest(ticker, method=varcov_var, confidences=CONFIDENCES)
    monte_carlo_backtest = do_backtest(ticker, method=monte_carlo_var, confidences=CONFIDENCES)

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2.3)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    document.add_heading('Value at Risk', 0)

    data = ticker.history(period='12mo')
    average_price = round(data['Close'].mean(), 2)

    basic_info = {
        'Název:': 'longName',
        'Symbol:': 'symbol',
        'Sektor:': 'sector',
        'Tržní kapitalizace:': 'marketCap',
        'P/E poměr:': 'forwardPE',
    }
    for key, value in list(basic_info.items()):
        if value in ticker.info:
            basic_info[key] = ticker.info[value]
            if isinstance(basic_info[key], (int, float)):
                basic_info[key] = '{:,}'.format(basic_info[key]).replace(',', ' ')
        else:
            del basic_info[key]
    basic_info['Průměrná cena (za rok)'] = f'{average_price}'
    insert_info(document, basic_info)

    wanted_fields = [
        'Total Assets',
        'Total Liabilities Net Minority Interest',
        'Total Equity Gross Minority Interest',
        'Total Capitalization',
        'Common Stock Equity',
        'Net Tangible Assets',
        'Working Capital',
        'Invested Capital',
        'Tangible Book Value',
        'Total Debt',
        'Net Debt',
        'Share Issued',
        'Ordinary Shares Number',
        'Treasury Shares Number'
    ]
    if not ticker.balance_sheet.empty:
        heading = document.add_heading('Balance Sheet [numbers in thousands]', 1)
        heading.paragraph_format.space_before = Pt(4)
        insert_table_filtered(document, ticker.balance_sheet.dropna(), wanted_fields)
        document.add_section(WD_SECTION.NEW_PAGE)

    wanted_fields3 = [
        'Total Revenue',
        'Cost of Revenue',
        'Gross Profit',
        'Operating Expense',
        'Operating Income',
        'Net Non Operating Interest Income Expense',
        'Other Income Expense',
        'Pretax Income',
        'Tax Provision',
        'Net Income Common Stockholders',
        'Diluted NI Available to Com Stockholders',
        # Zakomentářované řádky jsou většinou nulové
        #'Basic EPS',
        #'Diluted EPS',
        'Basic Average Shares',
        'Diluted Average Shares',
        'Total Operating Income as Reported',
        'Total Expenses',
        'Net Income from Continuing & Discontinued Operation',
        'Normalized Income',
        'Interest Income',
        'Interest Expense',
        'Net Interest Income',
        'EBIT',
        'EBITDA',
        'Reconciled Cost of Revenue',
        'Reconciled Depreciation',
        'Net Income from Continuing Operation Net Minority Interest',
        'Total Unusual Items Excluding Goodwill',
        'Total Unusual Items',
        'Normalized EBITDA',
        'Tax Rate for Calcs',
        'Tax Effect of Unusual Items'
    ]
    if not ticker.financials.empty:
        heading_income = document.add_heading('Income Statement [numbers in thousands]', 1)
        heading_income.paragraph_format.space_before = Pt(4)
        insert_table_filtered(document, ticker.financials.dropna(), wanted_fields3)
        document.add_section(WD_SECTION.NEW_PAGE)

    wanted_fields2 = [
        'Operating Cash Flow',
        'Investing Cash Flow',
        'Financing Cash Flow',
        'End Cash Position',
        'Capital Expenditure',
        'Issuance of Capital Stock',
        'Issuance of Debt',
        'Repayment of Debt',
        'Repurchase of Capital Stock',
        'Free Cash Flow',
        'Income Tax Paid Supplemental Data',
        'Interest Paid Supplemental Data'
    ]
    if not ticker.cashflow.empty:
        heading_cash_flow = document.add_heading('Cash Flow [numbers in thousands]', 1)
        heading_cash_flow.paragraph_format.space_before = Pt(4)
        insert_table_filtered(document, ticker.cashflow.dropna(), wanted_fields2)
        document.add_section(WD_SECTION.NEW_PAGE)
    
    try_insert_sustainability(document, ticker.ticker)

    document.add_heading('Historic Method', 1)
    add_graph(document, historic_graph)
    document.add_heading('Backtesting Historic VaR', 2)
    add_graph(document, historic_backtest)

    document.add_heading('Variance-Covariance Method', 1)
    add_graph(document, varcov_graph)
    document.add_heading('Backtesting Variance-Covariance VaR', 2)
    add_graph(document, varcov_backtest)

    document.add_heading('Monte Carlo Method', 1)
    add_graph(document, monte_carlo_graph)
    document.add_heading('Backtesting Monte Carlo VaR', 2)
    add_graph(document, monte_carlo_backtest)

    document.save(output_path)

# Create output directory
if not os.path.exists(OUTPUT_DIRECTORY):
    os.makedirs(OUTPUT_DIRECTORY)

# Expects a text file 'TICKERS.txt',
# where each line represents one yfinance ticker.

i = 0
with open('TICKERS.txt', 'r') as f:
    for line in f:
        i += 1
        stock = line.strip()
        if stock == '': continue
        ticker = yf.Ticker(stock)
        print(f"Generating report #{i} - {stock}")
        create_report(ticker, f'{OUTPUT_DIRECTORY}/{stock}.docx')
